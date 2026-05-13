"""
Document Processor: format routing, parsing, and multi-strategy chunking
"""
import os
import re
import logging
from typing import List, Dict, Optional
from enum import Enum

logger = logging.getLogger(__name__)


class ChunkStrategy(str, Enum):
    RECURSIVE = "recursive"
    HEADING = "heading"
    PARENT_CHILD = "parent_child"


class DocumentProcessor:
    """Format-aware document parser with multi-strategy chunking"""

    SUPPORTED_EXTENSIONS = {".pdf", ".txt", ".md", ".docx"}

    def __init__(
        self,
        strategy: str = "parent_child",
        chunk_size: int = 500,
        chunk_overlap: int = 75,
        parent_chunk_size: int = 1500,
    ):
        self.strategy = ChunkStrategy(strategy)
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        self.parent_chunk_size = parent_chunk_size

    def process(self, file_path: str) -> List[Dict]:
        """Full pipeline: parse -> chunk -> return chunks with metadata"""
        ext = os.path.splitext(file_path)[1].lower()
        if ext not in self.SUPPORTED_EXTENSIONS:
            raise ValueError(f"Unsupported file type: {ext}")

        text, extra_meta = self._parse(file_path, ext)

        if self.strategy == ChunkStrategy.PARENT_CHILD:
            chunks = self._parent_child_chunk(text)
        elif self.strategy == ChunkStrategy.HEADING:
            chunks = self._heading_aware_chunk(text)
        else:
            chunks = self._recursive_chunk(text)

        for c in chunks:
            c["metadata"].update(extra_meta)

        return chunks

    def _parse(self, file_path: str, ext: str):
        """Route to format-specific parser, return (text, extra_metadata)"""
        ext = ext.lower()
        if ext == ".pdf":
            return self._parse_pdf(file_path)
        elif ext == ".docx":
            return self._parse_docx(file_path)
        elif ext == ".md":
            return self._parse_md(file_path)
        else:
            return self._parse_txt(file_path)

    def _parse_pdf(self, file_path: str):
        """Parse PDF with structure awareness via PyMuPDF"""
        import fitz
        doc = fitz.open(file_path)
        pages_text = []
        metadata = {
            "doc_type": "pdf",
            "page_count": len(doc),
            "title": doc.metadata.get("title", os.path.basename(file_path)),
        }

        for page_num, page in enumerate(doc, start=1):
            blocks = page.get_text("dict")["blocks"]
            page_blocks = []
            for block in blocks:
                if block["type"] == 0:  # text
                    for line in block["lines"]:
                        text = "".join(span["text"] for span in line["spans"])
                        font_size = line["spans"][0]["size"] if line["spans"] else 12
                        is_bold = any(
                            "Bold" in s.get("font", "") for s in line["spans"]
                        )
                        page_blocks.append({
                            "text": text,
                            "font_size": font_size,
                            "is_bold": is_bold,
                        })
                elif block["type"] == 1:  # image
                    continue
            pages_text.append({"page_num": page_num, "blocks": page_blocks})

        lines = []
        for page in pages_text:
            prev_font = None
            for b in page["blocks"]:
                font = round(b["font_size"], 1)
                text = b["text"].strip()
                if not text:
                    continue
                if font and prev_font and font > prev_font + 2:
                    lines.append(f"\n## {text}")
                elif b["is_bold"] and len(text) < 80:
                    lines.append(f"\n### {text}")
                else:
                    lines.append(text)
                prev_font = font
            lines.append("")

        doc.close()
        return "\n".join(lines).strip(), metadata

    def _parse_docx(self, file_path: str):
        """Parse DOCX with style awareness via python-docx"""
        from docx import Document
        doc = Document(file_path)
        lines = []
        metadata = {
            "doc_type": "docx",
            "page_count": 0,
            "title": os.path.basename(file_path),
        }

        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue
            style_name = para.style.name if para.style else ""
            if "Heading 1" in style_name:
                lines.append(f"\n# {text}")
            elif "Heading 2" in style_name:
                lines.append(f"\n## {text}")
            elif "Heading 3" in style_name:
                lines.append(f"\n### {text}")
            elif "List" in style_name:
                lines.append(f"- {text}")
            else:
                lines.append(text)

        for table in doc.tables:
            lines.append("\n")
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                lines.append("| " + " | ".join(cells) + " |")

        return "\n".join(lines).strip(), metadata

    def _parse_md(self, file_path: str):
        """Parse Markdown directly"""
        import markdown
        with open(file_path, "r", encoding="utf-8") as f:
            text = f.read()
        metadata = {
            "doc_type": "md",
            "page_count": 0,
            "title": os.path.basename(file_path),
        }
        return text, metadata

    def _parse_txt(self, file_path: str):
        """Parse plain text"""
        with open(file_path, "r", encoding="utf-8") as f:
            text = f.read()
        metadata = {
            "doc_type": "txt",
            "page_count": 0,
            "title": os.path.basename(file_path),
        }
        return text, metadata

    def _recursive_chunk(self, text: str) -> List[Dict]:
        """Recursive character chunking with sentence boundary awareness"""
        chunks = []
        separators = ["\n\n", "\n", "。", "！", "？", "；", "\n"]

        text = self._split_sentences(text)
        paragraphs = [p.strip() for p in re.split(r'\n\s*\n', text) if p.strip()]

        current_chunk = ""
        for para in paragraphs:
            if len(current_chunk) + len(para) <= self.chunk_size:
                current_chunk += para + "\n"
            else:
                if current_chunk.strip():
                    chunks.append(current_chunk.strip())
                current_chunk = para + "\n"

        if current_chunk.strip():
            chunks.append(current_chunk.strip())

        result = []
        for i, chunk in enumerate(chunks):
            result.append({
                "text": chunk,
                "metadata": {
                    "chunk_index": i,
                    "total_chunks": len(chunks),
                    "char_count": len(chunk),
                    "chunk_id": f"chunk_{i}",
                    "strategy": "recursive",
                }
            })

        return result

    def _heading_aware_chunk(self, text: str) -> List[Dict]:
        """Chunk by markdown heading boundaries, preserving section hierarchy"""
        chunks = []
        sections = re.split(r'(\n#{1,3}\s+.*)', text)
        current_section = ""
        current_heading = ""
        section_path = []
        heading_level = 0

        for part in sections:
            heading_match = re.match(r'^(#{1,3})\s+(.*)', part.strip())
            if heading_match:
                if current_section.strip():
                    chunks.append({
                        "text": current_section.strip(),
                        "metadata": {
                            "section_title": current_heading,
                            "section_path": " > ".join(section_path) if section_path else current_heading,
                            "heading_level": heading_level,
                            "chunk_id": f"heading_{len(chunks)}",
                        }
                    })

                level = len(heading_match.group(1))
                title = heading_match.group(2).strip()
                section_path = section_path[:level - 1] + [title]
                current_heading = title
                heading_level = level
                current_section = part + "\n"
            else:
                if current_section:
                    sub_chunks = self._split_if_needed(part.strip(), current_heading)
                    for sc in sub_chunks:
                        chunks.append(sc)
                else:
                    current_section = part

        if current_section.strip():
            chunks.append({
                "text": current_section.strip(),
                "metadata": {
                    "section_title": current_heading,
                    "section_path": " > ".join(section_path) if section_path else current_heading,
                    "heading_level": heading_level,
                    "chunk_id": f"heading_{len(chunks)}",
                }
            })

        for i, c in enumerate(chunks):
            c["metadata"]["chunk_index"] = i
            c["metadata"]["total_chunks"] = len(chunks)
            c["metadata"]["char_count"] = len(c["text"])
            c["metadata"]["strategy"] = "heading"

        return chunks

    def _split_if_needed(self, text: str, section_title: str) -> List[Dict]:
        """Split long section content into smaller chunks"""
        if len(text) <= self.chunk_size:
            return [{
                "text": text,
                "metadata": {"section_title": section_title}
            }]

        chunks = []
        sentences = self._split_sentences(text)
        current = ""
        for sent in sentences:
            if len(current) + len(sent) <= self.chunk_size:
                current += sent
            else:
                if current.strip():
                    chunks.append({"text": current.strip(), "metadata": {"section_title": section_title}})
                current = sent
        if current.strip():
            chunks.append({"text": current.strip(), "metadata": {"section_title": section_title}})
        return chunks

    def _parent_child_chunk(self, text: str) -> List[Dict]:
        """Parent-child chunking: small children for retrieval, large parents for generation"""
        from collections import defaultdict

        paragraphs = [p.strip() for p in re.split(r'\n\s*\n', text) if p.strip()]

        parent_groups = defaultdict(list)
        current_parent_id = 0
        current_parent_size = 0

        for para in paragraphs:
            if current_parent_size + len(para) > self.parent_chunk_size and current_parent_size > 0:
                current_parent_id += 1
                current_parent_size = 0
            parent_groups[current_parent_id].append(para)
            current_parent_size += len(para)

        chunks = []
        child_index = 0
        for parent_id, parent_paras in parent_groups.items():
            parent_text = "\n".join(parent_paras)

            children = self._split_to_children(parent_text)
            for child in children:
                chunks.append({
                    "text": child,
                    "metadata": {
                        "chunk_id": f"child_{child_index}",
                        "chunk_index": child_index,
                        "parent_id": parent_id,
                        "parent_text": parent_text,
                        "role": "child",
                        "char_count": len(child),
                        "strategy": "parent_child",
                    }
                })
                child_index += 1

        total = len(chunks)
        for c in chunks:
            c["metadata"]["total_chunks"] = total

        return chunks

    def _split_to_children(self, text: str) -> List[str]:
        """Split parent text into overlapping child chunks"""
        sentences = self._split_sentences(text)
        if not sentences:
            return [text]

        chunks = []
        current = ""
        for sent in sentences:
            if len(current) + len(sent) <= self.chunk_size:
                current += sent
            else:
                if current.strip():
                    chunks.append(current.strip())
                current = self._get_overlap_tail(current, sent) + sent

        if current.strip():
            chunks.append(current.strip())

        return chunks if chunks else [text]

    def _get_overlap_tail(self, text: str, next_sent: str) -> str:
        """Extract overlapping tail from previous chunk"""
        if self.chunk_overlap <= 0:
            return ""
        overlap_chars = text[-self.chunk_overlap:] if len(text) > self.chunk_overlap else text
        last_break = max(
            overlap_chars.rfind("。"),
            overlap_chars.rfind("！"),
            overlap_chars.rfind("？"),
            overlap_chars.rfind("\n"),
        )
        if last_break > 0:
            return overlap_chars[last_break + 1:]
        return overlap_chars

    def _split_sentences(self, text: str) -> List[str]:
        """Split text into sentences with Chinese boundary awareness"""
        try:
            from sentencesplit import sentencebreak
            sentences = sentencebreak(text)
            return [s.strip() for s in sentences if s.strip()]
        except ImportError:
            sentences = re.split(r'(?<=[。！？.!?])\s*', text)
            return [s.strip() for s in sentences if s.strip()]
